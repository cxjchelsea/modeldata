# _*_ coding: utf-8 _*_
# @File:    doc2pdf
# @Time:    2024/5/7 9:31
# @Author:  chenxuejiao
# @Contact: 15801102378@163.com
# @Version: V 0.1

from docx import Document
from docx.oxml.ns import qn
import os
from docx2pdf import convert

def extract_paragraphs_by_title(docx_file, title):
    """
    Extracts paragraphs from a Word document by a specified title.

    Parameters:
        - docx_file (str): The path to the Word document.
        - title (str): The title of the section to extract paragraphs from.

    Returns:
        - list of str: The paragraphs belonging to the specified title.
    """
    paragraphs = []
    doc = Document(docx_file)
    found_title = False  # Variable to track if the title is found
    for paragraph in doc.paragraphs:
        if found_title and paragraph.style.name in ['Heading 1', 'Heading 2', 'Heading 3']:
            # Next title encountered, stop collecting paragraphs
            break
        if paragraph.style.name == 'Heading 3' and paragraph.text == title:
            # Found the title, start collecting paragraphs until next title
            found_title = True
            paragraphs.append(paragraph.text)
            continue
        if found_title:
            # Collect paragraphs after the found title
            paragraphs.append(paragraph.text)

    return paragraphs

def insert_paragraphs_to_new_doc(paragraphs, new_docx_file, chinese_font_name, english_font_name):
    """
    Inserts paragraphs into a new Word document with specified fonts for Chinese and English.

    Parameters:
        - paragraphs (list of str): The paragraphs to insert.
        - new_docx_file (str): The path to the new Word document.
        - chinese_font_name (str): The name of the font to use for Chinese characters.
        - english_font_name (str): The name of the font to use for English characters.
    """
    doc = Document()
    for paragraph_text in paragraphs:
        p = doc.add_paragraph()
        for char in paragraph_text:
            run = p.add_run(char)
            if ord(char) < 128:  # 英文字符的 Unicode 范围为 0-127
                run.font.name = english_font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), english_font_name)
            else:
                run.font.name = chinese_font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font_name)
    doc.save(new_docx_file)

def convert_docx_to_pdf(docx_file, title):
    """
    Converts a Word document to PDF.

    Parameters:
        - docx_file (str): The path to the Word document.
        - title (str): The title to use for naming the PDF file.
    """
    pdf_file = f"{title}.pdf"
    try:
        convert(docx_file, pdf_file)
        print(f"Conversion successful! PDF saved as '{pdf_file}'.")
    except Exception as e:
        print(f"Conversion failed: {e}")

def delete_docx(docx_file):
    """
    Deletes a Word document.

    Parameters:
        - docx_file (str): The path to the Word document.
    """
    try:
        os.remove(docx_file)
        print(f"Document '{docx_file}' has been deleted successfully.")
    except OSError as e:
        print(f"Error: {e.strerror}")

import pandas as pd
def extract_column_data(excel_file, sheet_name, column_name):
    """
    从 Excel 文件的特定工作表中提取特定列的数据并以数组形式返回。

    :param excel_file: Excel 文件路径
    :param sheet_name: 工作表名称
    :param column_name: 要提取的列的名称
    :return: 包含列数据的数组
    """
    # 读取 Excel 文件中的特定工作表
    df = pd.read_excel(excel_file, sheet_name=sheet_name)

    # 提取特定列的数据并转换为数组
    column_data = df[column_name].dropna().tolist()

    return column_data


# 示例使用
# excel_file = "complement.xlsx"  # Excel 文件路径
# sheet_name = "Sheet1"  # 工作表名称
# column_name = "Complement"  # 要提取的列的名称
#
# titles = extract_column_data(excel_file, sheet_name, column_name)
titles = ['冷连剪切轧轧件定尺长度计算模型']

docx_file = "轧制机理基础模型库模型说明文档.docx"  # 原始 Word 文档路径
new_docx_file_prefix = "output"  # 新的 Word 文档路径前缀
chinese_font_name = "宋体"  # 中文字体名称
english_font_name = "Arial"  # 英文字体名称

def process_title(docx_file, title, new_docx_file_prefix, chinese_font_name, english_font_name):
    paragraphs = extract_paragraphs_by_title(docx_file, title)
    new_docx_file = f"{new_docx_file_prefix}_{title}.docx"
    insert_paragraphs_to_new_doc(paragraphs, new_docx_file, chinese_font_name, english_font_name)
    convert_docx_to_pdf(new_docx_file, title)
    # delete_docx(new_docx_file)

for title in titles:
    process_title(docx_file, title, new_docx_file_prefix, chinese_font_name, english_font_name)

